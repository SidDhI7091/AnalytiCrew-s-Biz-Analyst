import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document

# Initialize Firebase Admin SDK
cred = credentials.Certificate(r"serviceAccountKey.json")
firebase_admin.initialize_app(cred)

# Reference to Firestore database
db = firestore.client()

# Fetch requirements from Firestore
def fetch_requirements_from_firestore(project_id):
    try:
        # Fetch project-specific requirements from Firestore
        requirements_ref = db.collection('projects').document(project_id).collection('extracted_requirements')
        requirements = requirements_ref.stream()

        # Extract the requirements into a list of dictionaries
        fetched_requirements = []
        for req in requirements:
            fetched_requirements.append(req.to_dict())
        return fetched_requirements
    except Exception as e:
        print(f"Error fetching data from Firebase: {e}")
        return []

# Function to replace placeholders in the Word document
def replace_placeholder(doc, placeholder, value):
    for p in doc.paragraphs:
        if placeholder in p.text:
            inline = p.runs
            for i in inline:
                if placeholder in i.text:
                    i.text = i.text.replace(placeholder, value)
    return doc

# Function to create the document with fetched requirements
def create_doc(project_id):
    # Fetch requirements from Firestore
    requirements = fetch_requirements_from_firestore(project_id)
    
    if not requirements:
        print("No requirements found or error fetching data.")
        return

    # Load the template
    try:
        doc = Document('template.docx')
    except FileNotFoundError:
        print("Error: The template file 'template.docx' was not found.")
        return

    # Define context or data for placeholders
    context = {
        'project_name': 'Sample Project',
        'document_title': 'Software Requirements Specification',
        'document_author': 'John Doe',
        'date': '2025-04-27',
        'version': '1.0',
        'extracted_requirements': requirements
    }

    # Replace placeholders with actual values
    doc = replace_placeholder(doc, '{{ project_name }}', context['project_name'])
    doc = replace_placeholder(doc, '{{ document_title }}', context['document_title'])
    doc = replace_placeholder(doc, '{{ document_author }}', context['document_author'])
    doc = replace_placeholder(doc, '{{ date }}', context['date'])
    doc = replace_placeholder(doc, '{{ version }}', context['version'])

    # Replace requirements
    for req in context['extracted_requirements']:
        doc.add_paragraph(f"- Requirement: {req['requirement_text']}")
        #doc.add_paragraph(f"  User: {req['user']}")
        #doc.add_paragraph(f"  Goal: {req['goal']}")
        doc.add_paragraph(f"  Category: {req['category']}")
        doc.add_paragraph(f"  Priority: {req['priority']}")

    # Save the updated document
    doc.save('output.docx')
    print("Document created successfully!")

# Example usage
project_id = "TRAN"  
create_doc(project_id)