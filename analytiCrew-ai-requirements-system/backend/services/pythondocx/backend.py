from flask import Flask, request, jsonify
import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
import os

app = Flask(__name__)

# Firebase initialization
cred = credentials.Certificate(r"serviceAccountKey.json")
firebase_admin.initialize_app(cred)
db = firestore.client()

# Fetch requirements from Firestore
def fetch_requirements_from_firestore(project_id):
    try:
        requirements_ref = db.collection('projects').document(project_id).collection('extracted_requirements')
        requirements = requirements_ref.stream()
        fetched_requirements = []
        for req in requirements:
            fetched_requirements.append(req.to_dict())
        return fetched_requirements
    except Exception as e:
        print(f"Error fetching data from Firebase: {e}")
        return []

# Function to replace placeholders in the Word document
def replace_placeholder(doc, placeholder, value):
    for p in doc.paragraphs:
        if placeholder in p.text:
            inline = p.runs
            for i in inline:
                if placeholder in i.text:
                    i.text = i.text.replace(placeholder, value)
    return doc

# API endpoint to trigger document creation
@app.route('/generate-doc', methods=['POST'])
def generate_doc():
    try:
        project_id = request.json.get('project_id')
        requirements = fetch_requirements_from_firestore(project_id)
        
        if not requirements:
            return jsonify({"error": "No requirements found or error fetching data."}), 400

        # Load the template
        try:
            doc = Document('template.docx')
        except FileNotFoundError:
            return jsonify({"error": "The template file 'template.docx' was not found."}), 400

        context = {
            'project_name': 'Sample Project',
            'document_title': 'Software Requirements Specification',
            'document_author': 'John Doe',
            'date': '2025-04-27',
            'version': '1.0',
            'extracted_requirements': requirements
        }

        # Replace placeholders with actual values
        doc = replace_placeholder(doc, '{{ project_name }}', context['project_name'])
        doc = replace_placeholder(doc, '{{ document_title }}', context['document_title'])
        doc = replace_placeholder(doc, '{{ document_author }}', context['document_author'])
        doc = replace_placeholder(doc, '{{ date }}', context['date'])
        doc = replace_placeholder(doc, '{{ version }}', context['version'])

        # Replace requirements
        for req in context['extracted_requirements']:
            doc.add_paragraph(f"- Requirement: {req['requirement_text']}")
            doc.add_paragraph(f"  Category: {req['category']}")
            doc.add_paragraph(f"  Priority: {req['priority']}")

        # Save the updated document
        output_path = 'static/output.docx'
        doc.save(output_path)
        
        return jsonify({"message": "Document created successfully!", "file_path": output_path}), 200

    except Exception as e:
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(debug=True)
